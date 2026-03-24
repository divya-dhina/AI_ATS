from docx import Document

def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        full_text = []

        # 1️⃣ Normal paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # 2️⃣ Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)

        # 3️⃣ Headers & Footers (VERY IMPORTANT)
        for section in doc.sections:
            header = section.header
            footer = section.footer

            for para in header.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)

            for para in footer.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)

        # 4️⃣ Deep XML extraction (handles textboxes/shapes)
        for element in doc.element.body.iter():
            if element.text and element.text.strip():
                full_text.append(element.text)

        text = "\n".join(full_text)
        print(f"DOCX extracted length: {len(text)}")

        return text.strip()

    except Exception as e:
        print(f"Error reading DOCX file {file_path}: {e}")
        return ""